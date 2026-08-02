# -*- coding: utf-8 -*-
"""
生成《情绪手账小程序优化规划文档》Word 文件
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def set_cell_border(cell, **kwargs):
    """为表格单元格设置边框（仅上/下/左/右）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                element = tcPr.find(qn(tag))
                if element is None:
                    element = docx.oxml.OxmlElement(tag)
                    tcPr.append(element)
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), edge_data.get('color', '000000'))


def add_heading_custom(doc, text, level=1):
    """添加标题并设置中文字体"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 1:
            run.font.color.rgb = RGBColor(0xE8, 0x85, 0x6A)
            run.font.size = Pt(20)
            run.font.bold = True
        elif level == 2:
            run.font.color.rgb = RGBColor(0x4A, 0x3A, 0x32)
            run.font.size = Pt(16)
            run.font.bold = True
        else:
            run.font.color.rgb = RGBColor(0x5D, 0x4A, 0x40)
            run.font.size = Pt(13)
            run.font.bold = True
    return heading


def add_paragraph_custom(doc, text, bold=False, size=11, color=None, align=None, first_line_indent=0):
    """添加正文段落并设置中文字体"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Cm(first_line_indent)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_bullet(doc, text, level=0):
    """添加项目符号段落"""
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x4A, 0x3A, 0x32)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    return p


def add_numbered(doc, text):
    """添加编号段落"""
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x4A, 0x3A, 0x32)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    return p


def add_feature_card(doc, title, tag, desc, interaction, value, difficulty):
    """以表格形式展示单个功能卡片"""
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Cm(3.5)
    table.columns[1].width = Cm(12.5)

    cells = [
        ('功能名称', title),
        ('功能定位', tag),
        ('功能描述', desc),
        ('交互方式', interaction),
        ('预期价值 / 难度', f"{value} ｜ 实现难度：{difficulty}"),
    ]
    for i, (k, v) in enumerate(cells):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(10.5)
            cell.vertical_alignment = 1  # CENTER
    # 加粗第一列
    for row in table.rows:
        for run in row.cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xE8, 0x85, 0x6A)
    doc.add_paragraph()


def main():
    doc = Document()

    # 默认正文字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x4A, 0x3A, 0x32)

    # 页面边距
    sections = doc.sections[0]
    sections.top_margin = Cm(2.5)
    sections.bottom_margin = Cm(2.5)
    sections.left_margin = Cm(2.5)
    sections.right_margin = Cm(2.5)

    # ==================== 封面 ====================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('情绪手账小程序')
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xE8, 0x85, 0x6A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('天气 × 心情 优化升级规划文档')
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x4A, 0x3A, 0x32)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('版本：V2.0 优化规划\n日期：2026 年 8 月\n产品形态：微信小程序')
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x8B, 0x77, 0x6B)

    doc.add_page_break()

    # ==================== 目录（手动） ====================
    add_heading_custom(doc, '目录', level=1)
    toc_items = [
        '一、项目概述与现状分析',
        '二、新增功能规划',
        '2.1 天气-心情趋势图（多维联动）',
        '2.2 动态心情天气壁纸',
        '2.3 心情日历热力图',
        '2.4 天气情绪共振指数',
        '2.5 情境化记录引导升级',
        '2.6 每周天气心情故事',
        '三、现有界面优化方向',
        '3.1 天气展示界面（首页沉浸式天气卡）',
        '3.2 心情记录界面',
        '3.3 历史记录 / 情绪日历界面',
        '3.4 晴雨表 / 洞察界面',
        '3.5 天气心情地图界面',
        '3.6 整体视觉与动效统一',
        '四、数据埋点与算法支撑',
        '五、实施里程碑与优先级',
        '六、预期效果与衡量指标',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0 if item.startswith('一、') or item.startswith('二、') or item.startswith('三、') or item.startswith('四、') or item.startswith('五、') or item.startswith('六、') else 0.8)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x4A, 0x3A, 0x32)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    doc.add_page_break()

    # ==================== 一、项目概述 ====================
    add_heading_custom(doc, '一、项目概述与现状分析', level=1)
    add_paragraph_custom(doc,
        '「情绪手账」是一款将实时天气与心情记录深度融合的小程序，用户可在记录情绪的同时，自动关联当日天气数据，逐步积累个人专属的天气-心情档案。'
        '当前版本已具备天气动态背景、沉浸式天气卡、心情网格选择、天气影响追问、情绪健康指数（EHI）、天气心情地图、情绪日历等核心能力，'
        '视觉风格以暖橘、米白、柔和阴影为主，整体氛围温馨治愈。')

    add_heading_custom(doc, '1.1 现状功能清单', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '模块'
    hdr_cells[1].text = '现有能力'
    hdr_cells[2].text = '待提升空间'
    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rows = [
        ('首页', '动态天气背景、沉浸式天气卡、心情网格、天气影响追问、今日心情轨迹、每日一句', '信息层次可更清晰，动效与数据关联可进一步可视化'),
        ('记录页', '日期时间、心情选择、天气快照、天气影响、身体感受、诱因、场景、多媒体', '填写路径较长，可引入智能引导和进度感知'),
        ('情绪日历', '月度打卡、颜色标记、详情弹窗', '缺乏天气-心情叠加层与热力趋势'),
        ('晴雨表', 'EHI 指数、情绪走势柱状图、天气关联柱状图', '图表维度单一，缺少趋势线与多天气对比'),
        ('天气心情地图', '二维分布表格、个性化洞察', '表格可读性有限，可升级为卡片故事化呈现'),
        ('我的', '个人资料、收藏、连续打卡', '可加入数据导出与隐私管理'),
    ]
    for r in rows:
        row_cells = table.add_row().cells
        for i, text in enumerate(r):
            row_cells[i].text = text
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(10.5)
    doc.add_paragraph()

    add_heading_custom(doc, '1.2 核心设计原则', level=2)
    add_bullet(doc, '温暖陪伴：保留暖橘、米白、毛玻璃等治愈元素，避免冷冰冰的数据感。')
    add_bullet(doc, '直观关联：让用户一眼看懂「天气如何影响我的心情」。')
    add_bullet(doc, '轻量可执行：优先做小程序能稳定承载、上线成本低的功能。')
    add_bullet(doc, '渐进呈现：数据量不足时给出温柔提示，数据积累后自动解锁深度洞察。')
    doc.add_page_break()

    # ==================== 二、新增功能规划 ====================
    add_heading_custom(doc, '二、新增功能规划', level=1)
    add_paragraph_custom(doc,
        '本章节围绕「天气 × 心情」的关联感知，提出 6 项创新功能。每项功能均包含功能描述、交互方式、预期价值与实现难度评估。')

    add_feature_card(
        doc,
        '天气-心情趋势图（多维联动）',
        '核心数据可视化',
        '在「晴雨表」页新增双轴趋势图：横轴为时间（日/周/月），左侧纵轴为心情评分（1-5 分），右侧纵轴为天气指标（温度、湿度、空气质量）。'
        '曲线与柱状叠加，用户可切换「温度 / 湿度 / 天气类型 / 空气质量」维度，直观看到心情波动与天气变化是否同步。'
        '数据不足时显示「再多记录 3 天即可解锁趋势图」的鼓励文案。',
        '点击图例切换维度；长按数据点查看当天记录摘要；双指捏合缩放时间轴。',
        '高：显著提升数据价值感，让用户从「记录」走向「理解自己」',
        '中（可使用 echarts-for-weixin 或自研 Canvas 折线图）'
    )

    add_feature_card(
        doc,
        '动态心情天气壁纸',
        '情感化首页背景',
        '首页背景由「纯天气动效」升级为「天气 + 心情」混合壁纸。例如：晴天+开心时背景是明亮的金色粒子与微笑云朵；'
        '雨天+低落时背景是灰蓝雨滴与缓慢下沉的水滴，但会有一束暖光逐渐亮起，传递「阴霾会过去」的关怀。'
        '壁纸根据当日最新记录的心情实时变化，每次进入首页都有微妙的情绪呼应。',
        '自动读取最近一次心情记录与实时天气，生成对应壁纸；支持用户在设置中关闭动态壁纸以节省电量。',
        '高：强化产品差异化与情感连接',
        '中（复用现有 CSS 动画，按心情叠加颜色/粒子层）'
    )

    add_feature_card(
        doc,
        '心情日历热力图',
        '时间维度情绪密度',
        '在「情绪日历」基础上，新增「心情热力图层」。每个日期格子的背景色深浅表示当天平均心情高低（暖橘色越深表示越积极，浅灰表示无记录）。'
        '点击日期后，底部弹出当日天气-心情简报：天气图标、心情均分、高频诱因、代表句摘录。',
        '月份切换时热力图以淡入动画呈现；点击日期展开详情；支持左右滑动切换月份。',
        '高：快速识别情绪周期与天气规律',
        '低（复用现有日历网格，增加颜色计算逻辑）'
    )

    add_feature_card(
        doc,
        '天气情绪共振指数',
        '量化关联强度',
        '基于历史数据计算「我的天气敏感度」指数（0-100）。例如：雨天心情明显偏低则「雨感指数」高；晴天心情普遍好则「阳光增益指数」高。'
        '指数以仪表盘形式展示，并配以一句解读：「雨天容易让你陷入沉思，记得多给自己一点温柔」。',
        '每周一首次打开小程序时弹出「本周天气情绪共振」卡片；指数页支持查看历史变化。',
        '中：把抽象关联转化为可分享的量化标签',
        '中（需要后台统计与前端仪表盘组件）'
    )

    add_feature_card(
        doc,
        '情境化记录引导升级',
        '降低记录成本',
        '记录页顶部根据当前天气和时段，动态生成一句引导语。例如：「下雨天，适合听雨声，也适合把心情写下来」。'
        '同时根据用户历史数据，智能推荐最可能的 3 个心情标签和 2 个诱因标签，减少选择路径。',
        '首页点击心情后进入记录页，自动预选对应心情；推荐标签可一键点选，也可继续自定义。',
        '高：降低记录门槛，提升日活',
        '低（基于规则 + 历史统计，无需复杂模型）'
    )

    add_feature_card(
        doc,
        '每周天气心情故事',
        '故事化数据回顾',
        '每周日晚生成一份「我的天气心情周报」：以叙事方式总结本周天气分布、心情高光时刻、最低谷那天、最特别的天气-心情组合。'
        '报告支持生成精美长图分享给好友，或保存到相册。',
        '在「晴雨表」页新增「周报」入口；每周首次进入弹出提示；一键生成长图。',
        '高：促进分享传播与用户回流',
        '中（报告模板 + Canvas/图片合成）'
    )

    doc.add_page_break()

    # ==================== 三、现有界面优化 ====================
    add_heading_custom(doc, '三、现有界面优化方向', level=1)
    add_paragraph_custom(doc,
        '在新增功能的同时，对现有页面进行系统性的视觉与交互优化，使整体体验更统一、更精致、更具情感化。')

    add_heading_custom(doc, '3.1 天气展示界面（首页沉浸式天气卡）', level=2)
    add_paragraph_custom(doc, '优化目标：让天气信息成为首页的视觉焦点，同时不抢夺心情记录的入口地位。', bold=True)
    add_bullet(doc, '视觉层次：将温度数字放大至 96rpx 并采用细体，城市与天气类型以胶囊标签形式置于左上角，减少文字堆叠。')
    add_bullet(doc, '关怀文案：根据天气+时段+最近一次心情，生成三段式文案（天气事实 + 情绪共鸣 + 行动建议），控制在 2 行以内。')
    add_bullet(doc, '实时标签：「实时天气」标签改为脉搏动画的绿色圆点，强化「正在同步」的感知。')
    add_bullet(doc, '穿衣建议：升级为图标+简短卡片的组合，雨天自动显示雨伞/雨衣建议，高温显示防暑提示。')
    add_bullet(doc, '动效节奏：背景云朵漂浮速度根据天气调整（晴天更慢、大风更快），雨天增加玻璃雨滴沿卡片边缘滑落的细节。')

    add_heading_custom(doc, '3.2 心情记录界面', level=2)
    add_paragraph_custom(doc, '优化目标：减少填写压力，让记录过程像聊天一样自然。', bold=True)
    add_bullet(doc, '步骤可视化：在顶部增加 3 步进度条（选心情 → 写感受 → 加标签），当前步骤高亮，已完成的步骤显示对勾。')
    add_bullet(doc, '心情选择：保留 4 列网格，但为每个心情增加微动效（选中时放大 1.1 倍并上浮 8rpx），未选项轻微降透明度，强化选择反馈。')
    add_bullet(doc, '智能推荐区：在「天气对我影响」「身体感受」「情绪诱因」上方增加「根据你以往记录，推荐…」的浅灰提示条。')
    add_bullet(doc, '多媒体入口：将图片与录音从底部平铺改为「+」展开浮层，默认收起，降低界面复杂度。')
    add_bullet(doc, '保存反馈：保存成功后展示 1.2 秒的庆祝动画（如心情 emoji 弹跳 + 粒子扩散），再返回首页。')

    add_heading_custom(doc, '3.3 历史记录 / 情绪日历界面', level=2)
    add_paragraph_custom(doc, '优化目标：让历史数据既有美感，又能快速洞察规律。', bold=True)
    add_bullet(doc, '月份切换：当前左右箭头改为可左右滑动的月份条，点击月份快速跳转，减少操作步数。')
    add_bullet(doc, '日期格：有记录的日期显示心情 emoji 与天气图标小徽章，空日期显示淡灰色占位圆点，保持视觉整齐。')
    add_bullet(doc, '详情弹窗：从底部中间弹窗改为底部 Sheet，背景模糊，增加「编辑」「分享日签」「删除」三个操作按钮。')
    add_bullet(doc, '连续打卡：在日历顶部增加火焰 icon + 连续天数，并展示「还差 X 天解锁本周成就」的轻激励。')

    add_heading_custom(doc, '3.4 晴雨表 / 洞察界面', level=2)
    add_paragraph_custom(doc, '优化目标：把零散图表升级为有叙事逻辑的洞察中心。', bold=True)
    add_bullet(doc, 'EHI 指数卡：采用环形进度条替代横向进度条，中心显示分数，外圈颜色随分数变化（绿-黄-橙-红）。')
    add_bullet(doc, '情绪走势：柱状图改为「柱状 + 7 日移动平均线」，帮助用户识别短期趋势。')
    add_bullet(doc, '天气关联：由单一柱状图升级为雷达图/玫瑰图，展示不同天气类型下的平均心情、记录频次、情绪波动幅度。')
    add_bullet(doc, '异常提醒：连续低谷天数 ≥ 3 天时，卡片颜色变为暖橙关怀色，文案避免医学诊断，仅提供温柔建议与资源入口。')

    add_heading_custom(doc, '3.5 天气心情地图界面', level=2)
    add_paragraph_custom(doc, '优化目标：把二维表格转化为可浏览、可探索的卡片流。', bold=True)
    add_bullet(doc, '卡片流布局：每个天气类型一张卡片，顶部是天气 emoji + 名称 + 记录次数，中间是心情分布圆环，底部是典型感受文案。')
    add_bullet(doc, '空状态：当某种天气记录不足 3 次时，卡片显示「再多记录几次，就能解锁这份天气的心情画像」。')
    add_bullet(doc, '筛选器：顶部增加「按心情筛选」（只看开心的日子 / 只看低落的日子），反向发现天气与情绪的隐藏关联。')

    add_heading_custom(doc, '3.6 整体视觉与动效统一', level=2)
    add_bullet(doc, '色彩系统：在现有暖橘主色基础上，为不同天气定义更鲜明的辅助色板（晴：金橘；雨：雾蓝；雪：冰青；夜：深紫）。')
    add_bullet(doc, '圆角与阴影：统一卡片圆角为 28rpx，阴影采用向上偏移 8rpx、模糊 28rpx、暖色调低透明度，保持柔和漂浮感。')
    add_bullet(doc, '图标风格：统一使用圆润线面结合的 emoji + 线性图标，避免不同页面图标风格割裂。')
    add_bullet(doc, '转场动画：页面切换增加 120ms 的淡入上滑；卡片展开使用 spring 弹性曲线，避免生硬。')
    add_bullet(doc, '无障碍：所有心情标签与天气图标增加语义化标签，确保读屏软件可识别。')

    doc.add_page_break()

    # ==================== 四、数据埋点与算法支撑 ====================
    add_heading_custom(doc, '四、数据埋点与算法支撑', level=1)
    add_paragraph_custom(doc, '为确保新增功能可衡量、可迭代，需同步完善数据体系。')

    add_heading_custom(doc, '4.1 关键埋点', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '事件'
    hdr[1].text = '指标'
    hdr[2].text = '用途'
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    events = [
        ('记录完成', '完成率、平均耗时、放弃率', '优化记录流程与引导'),
        ('趋势图查看', '查看次数、切换维度次数、停留时长', '评估数据可视化价值'),
        ('周报分享', '生成次数、分享次数、回流率', '衡量社交传播效果'),
        ('动态壁纸展示', '展示次数、关闭率', '平衡情感化与性能'),
        ('日历热力图点击', '点击次数、查看详情转化率', '优化历史回顾路径'),
    ]
    for e in events:
        row = table.add_row().cells
        for i, text in enumerate(e):
            row[i].text = text
            for p in row[i].paragraphs:
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(10.5)
    doc.add_paragraph()

    add_heading_custom(doc, '4.2 算法/规则支撑', level=2)
    add_bullet(doc, '天气-心情相关性计算：采用皮尔逊相关系数量化天气指标（温度、湿度、AQI）与心情评分的线性关系，仅当样本量 ≥ 10 时展示结论。')
    add_bullet(doc, '情绪健康指数（EHI）：沿用现有公式，新增「连续低谷天数」权重与天气波动修正因子。')
    add_bullet(doc, '智能推荐：基于最近 30 天记录频次与天气条件，使用简单规则引擎推荐心情与诱因，避免冷启动。')
    add_bullet(doc, '周报生成：采用模板填充 + 数据摘要，突出最高/最低分日期、最频繁天气、最特别组合。')

    doc.add_page_break()

    # ==================== 五、实施里程碑 ====================
    add_heading_custom(doc, '五、实施里程碑与优先级', level=1)
    add_paragraph_custom(doc, '按照「先基础体验、再数据可视化、最后社交传播」的节奏，分三个阶段推进。')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '阶段'
    hdr[1].text = '周期'
    hdr[2].text = '核心内容'
    hdr[3].text = '交付物'
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    milestones = [
        ('第一阶段：体验打磨', '2 周', '优化天气卡、记录页步骤可视化、心情选择动效、保存反馈、日历详情 Sheet', '首页/记录页/日历页高保真设计稿 + 前端实现'),
        ('第二阶段：数据可视化', '3 周', '天气-心情趋势图、心情日历热力图、天气情绪共振指数、晴雨表雷达图', '图表组件 + 洞察页重构'),
        ('第三阶段：情感化与传播', '2 周', '动态心情天气壁纸、情境化引导、每周天气心情故事与分享长图', '壁纸动效 + 周报生成与分享'),
    ]
    for m in milestones:
        row = table.add_row().cells
        for i, text in enumerate(m):
            row[i].text = text
            for p in row[i].paragraphs:
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(10.5)
    doc.add_paragraph()

    add_paragraph_custom(doc, '优先级建议：', bold=True)
    add_numbered(doc, 'P0：记录页步骤可视化、天气-心情趋势图、心情日历热力图（直接影响核心留存与数据价值）')
    add_numbered(doc, 'P1：动态心情天气壁纸、天气情绪共振指数、晴雨表雷达图（提升差异化与情感化）')
    add_numbered(doc, 'P2：每周天气心情故事、情境化智能推荐、分享长图（提升传播与活跃度）')

    doc.add_page_break()

    # ==================== 六、预期效果 ====================
    add_heading_custom(doc, '六、预期效果与衡量指标', level=1)
    add_paragraph_custom(doc, '优化上线后，建议从用户参与度、数据价值感、情感认同三个维度评估效果。')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '维度'
    hdr[1].text = '指标'
    hdr[2].text = '当前基线（预估）'
    hdr[3].text = '目标提升'
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    metrics = [
        ('参与度', '日活跃用户数（DAU）', '—', '+15%'),
        ('参与度', '人均每周记录次数', '—', '+25%'),
        ('数据价值感', '洞察页人均停留时长', '—', '+40%'),
        ('数据价值感', '趋势图/热力图查看率', '—', '≥ 35%'),
        ('情感认同', '动态壁纸关闭率', '—', '≤ 15%'),
        ('传播', '周报分享率', '—', '≥ 8%'),
        ('留存', '7 日留存率', '—', '+10%'),
    ]
    for m in metrics:
        row = table.add_row().cells
        for i, text in enumerate(m):
            row[i].text = text
            for p in row[i].paragraphs:
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(10.5)
    doc.add_paragraph()

    add_paragraph_custom(doc,
        '以上目标需在产品上线后通过 A/B 测试与埋点数据持续验证，建议每两周进行一次数据复盘，'
        '根据用户反馈调整功能优先级与视觉细节。')

    # 保存
    output_path = r'C:\Users\tflsy\Documents\Codex\2026-07-21\documents-plugin-documents-openai-primary-runtime\mood-journal-miniprogram\情绪手账_天气心情优化规划.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')


if __name__ == '__main__':
    main()
